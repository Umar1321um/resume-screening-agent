"""
parser.py
---------
Extracts raw text from resume files in PDF, DOCX, or plain-text format.

Design notes:
- Each format has a dedicated extractor function so failures are isolated
  and easy to debug (a bad PDF won't take down DOCX parsing, etc).
- Text is lightly normalized (whitespace collapsed) so downstream NLP
  (TF-IDF / regex) behaves consistently regardless of source format.
"""

from __future__ import annotations

import re
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text(file_path: str | Path) -> str:
    """Extract raw text from a resume file based on its extension."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        text = _extract_pdf(path)
    elif ext == ".docx":
        text = _extract_docx(path)
    elif ext == ".txt":
        text = _extract_txt(path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}' for {path.name}. "
            f"Supported types: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    return _normalize(text)


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    chunks = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            chunks.append(page_text)
    return "\n".join(chunks)


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]

    # Also pull text out of any tables (some resumes use table layouts).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _normalize(text: str) -> str:
    # Collapse excessive blank lines/whitespace but keep line breaks
    # (useful for section-based regex extraction later).
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_resumes(folder: str | Path) -> list[dict]:
    """Load and parse every supported resume file in a folder.

    Returns a list of dicts: {"file_name": ..., "file_path": ..., "text": ...}
    Files that fail to parse are skipped with a printed warning rather than
    crashing the whole batch run.
    """
    folder = Path(folder)
    results = []
    for path in sorted(folder.iterdir()):
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            text = extract_text(path)
            if not text.strip():
                print(f"[parser] Warning: {path.name} produced no extractable text.")
                continue
            results.append({
                "file_name": path.name,
                "file_path": str(path),
                "text": text,
            })
        except Exception as exc:  # noqa: BLE001 - isolate per-file failures
            print(f"[parser] Failed to parse {path.name}: {exc}")
    return results
